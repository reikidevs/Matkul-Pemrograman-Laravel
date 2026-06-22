# -*- coding: utf-8 -*-
"""
Generator Laporan Projek 2 - Sistem KostKu
Sesuai ketentuan: Times New Roman 12, spasi 1.5, margin 3 cm, output .docx
Berisi: Cover + BAB 4 (Hasil + screenshot + deskripsi)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SHOT = os.path.join(BASE, "docs", "screenshots")
OUT = os.path.join(BASE, "LAPORAN_PROJEK2_KOSTKU.docx")

FONT = "Times New Roman"
SIZE = 12

# ---------- Data Identitas (disamakan dengan Projek 01) ----------
IDENTITAS = {
    "judul": 'LAPORAN PROJEK 2\nIMPLEMENTASI SISTEM\nAPLIKASI MANAJEMEN KOST "KOSTKU"',
    "kelompok": "Kelompok Contoh Lengkap",
    "anggota": [
        ("Ahmad Rizki Kurniawan", "G.131.24.0001"),
        ("Siti Nurhaliza", "G.131.24.0002"),
        ("Budi Santoso", "G.131.24.0003"),
        ("Dewi Lestari", "G.131.24.0004"),
        ("Eko Prasetyo", "G.131.24.0005"),
        ("Fitri Handayani", "G.131.24.0006"),
    ],
    "matkul": "Pemrograman Berbasis Kerangka Kerja (PBKK)",
    "dosen": "Ahmad Rifa'i, S.Kom., M.Kom.",
    "prodi": "PROGRAM STUDI SISTEM INFORMASI",
    "fakultas": "FAKULTAS TEKNOLOGI INFORMASI DAN KOMUNIKASI",
    "univ": "UNIVERSITAS SEMARANG",
    "tahun": "2026",
}


def set_base_style(doc):
    """Set font default Times New Roman 12, spasi 1.5."""
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def set_margins(section):
    """Margin 3 cm semua sisi."""
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)


def add_para(doc, text="", bold=False, align=None, size=SIZE, space_after=0,
             italic=False, line_spacing=1.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    return p


print("Script siap. Bagian fungsi inti selesai dibuat.")


def build_cover(doc):
    """Membuat halaman Cover (disamakan dengan Projek 01)."""
    # Judul
    for line in IDENTITAS["judul"].split("\n"):
        add_para(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, space_after=0)
    add_para(doc, "", space_after=18)

    # Logo placeholder
    logo = os.path.join(SHOT, "logo.png")
    if os.path.exists(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo, width=Cm(4))
    else:
        add_para(doc, "[ LOGO UNIVERSITAS SEMARANG ]", align=WD_ALIGN_PARAGRAPH.CENTER,
                 italic=True, space_after=6)
    add_para(doc, "", space_after=18)

    # Disusun oleh
    add_para(doc, "Disusun oleh:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, IDENTITAS["kelompok"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for nama, nim in IDENTITAS["anggota"]:
        add_para(doc, f"{nama}  -  {nim}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", space_after=18)

    # Mata Kuliah
    add_para(doc, "Mata Kuliah:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, IDENTITAS["matkul"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Dosen Pengampu
    add_para(doc, "Dosen Pengampu:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, IDENTITAS["dosen"], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Institusi
    add_para(doc, IDENTITAS["prodi"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, IDENTITAS["fakultas"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, IDENTITAS["univ"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, IDENTITAS["tahun"], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_screenshot(doc, filename, caption_no, caption_text):
    """Sisipkan screenshot bila ada, atau kotak placeholder bila tidak."""
    path = None
    for ext in (".png", ".jpg", ".jpeg"):
        cand = os.path.join(SHOT, filename + ext)
        if os.path.exists(cand):
            path = cand
            break

    if path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Lebar maksimal area cetak = 21cm - 3cm - 3cm = 15cm
        p.add_run().add_picture(path, width=Cm(14))
        p.paragraph_format.space_before = Pt(6)
    else:
        # Placeholder box
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[ Tempelkan screenshot di sini: {filename}.png ]")
        run.font.name = FONT
        run.font.size = Pt(SIZE)
        run.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # border kotak sederhana via shading paragraf
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "dashed")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "10")
            el.set(qn("w:color"), "AAAAAA")
            borders.append(el)
        pPr.append(borders)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(18)

    # Caption
    cap = add_para(doc, f"Gambar 4.{caption_no} {caption_text}",
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=12)
    cap.runs[0].italic = True


# ---------- Konten BAB 4: (file, judul subbab, deskripsi) ----------
HASIL = [
    ("01_login", "Halaman Login",
     "Halaman login merupakan halaman pertama yang muncul saat aplikasi diakses. "
     "Pengguna harus memasukkan email dan password yang valid untuk dapat masuk ke "
     "sistem. Sistem menggunakan autentikasi bawaan Laravel dengan validasi kredensial. "
     "Apabila email atau password salah, sistem menampilkan pesan kesalahan. Terdapat "
     "tiga peran pengguna, yaitu Admin, Pemilik Kost, dan Penghuni, yang masing-masing "
     "memiliki hak akses berbeda setelah berhasil login."),

    ("02_dashboard_pemilik", "Dashboard Pemilik Kost",
     "Setelah pemilik kost berhasil login, sistem menampilkan halaman dashboard yang "
     "berisi ringkasan informasi penting. Dashboard menampilkan statistik berupa jumlah "
     "total kamar, jumlah penghuni, tingkat hunian (occupancy rate), serta total "
     "pendapatan. Selain itu, terdapat pula notifikasi mengenai pembayaran yang menunggu "
     "konfirmasi dan komplain yang masih terbuka, sehingga pemilik dapat segera "
     "menindaklanjutinya."),

    ("03_kelola_kamar", "Halaman Kelola Kamar",
     "Halaman Kelola Kamar menampilkan daftar seluruh kamar kost dalam bentuk tabel. "
     "Setiap kamar menampilkan informasi nomor kamar, tipe, harga sewa, serta status "
     "kamar (tersedia, terisi, atau dalam perawatan). Pemilik dapat menambah, mengubah, "
     "dan menghapus data kamar melalui tombol aksi yang tersedia pada setiap baris."),

    ("04_form_kamar", "Form Tambah/Edit Kamar",
     "Form ini digunakan untuk menambah kamar baru atau mengubah data kamar yang sudah "
     "ada. Form berisi input nomor kamar, tipe kamar, harga sewa, luas, fasilitas, dan "
     "status kamar. Sistem melakukan validasi terhadap data yang dimasukkan untuk "
     "memastikan kelengkapan dan kebenaran data sebelum disimpan ke basis data."),

    ("05_kelola_penghuni", "Halaman Kelola Penghuni",
     "Halaman Kelola Penghuni menampilkan daftar penghuni kost beserta status "
     "pendaftarannya. Pemilik dapat menyetujui (approve) atau menolak (reject) "
     "pendaftaran penghuni baru. Data penghuni mencakup nama, kontak, dan kamar yang "
     "ditempati. Fitur ini mendukung proses verifikasi calon penghuni sebelum resmi "
     "menempati kamar."),

    ("06_konfirmasi_pembayaran", "Halaman Konfirmasi Pembayaran",
     "Halaman ini menampilkan seluruh data pembayaran sewa yang dilakukan penghuni. "
     "Setiap pembayaran menampilkan nama penghuni, kamar, periode, jumlah, tanggal "
     "transfer, serta status pembayaran (pending, approved, atau rejected). Pemilik "
     "dapat menyetujui atau menolak pembayaran setelah melakukan verifikasi bukti "
     "transfer yang diunggah penghuni."),

    ("07_lihat_bukti_transfer", "Tampilan Bukti Transfer",
     "Ketika pemilik menekan tombol \"Lihat Bukti\", sistem menampilkan gambar bukti "
     "transfer yang diunggah penghuni dalam tampilan layar penuh (modal). Fitur ini "
     "memudahkan pemilik untuk memeriksa keabsahan bukti pembayaran sebelum memberikan "
     "persetujuan. Apabila pembayaran tidak valid, pemilik dapat menolak dengan "
     "menyertakan alasan penolakan."),

    ("08_kelola_komplain", "Halaman Kelola Komplain",
     "Halaman Kelola Komplain menampilkan daftar komplain yang diajukan penghuni. Setiap "
     "komplain memiliki kategori, judul, deskripsi, dan status penanganan. Pemilik dapat "
     "memperbarui status komplain mulai dari terbuka (open), sedang ditangani (in "
     "progress), hingga selesai (resolved), sehingga proses penanganan keluhan menjadi "
     "lebih transparan dan terdokumentasi."),

    ("09_dashboard_penghuni", "Dashboard Penghuni",
     "Dashboard penghuni menampilkan informasi ringkas mengenai kamar yang ditempati, "
     "tagihan terkini, serta status komplain yang sedang diajukan. Tampilan ini membantu "
     "penghuni memantau kewajiban pembayaran dan informasi kost secara cepat dalam satu "
     "halaman."),

    ("10_kamar_saya", "Halaman Kamar Saya",
     "Halaman Kamar Saya menampilkan detail kamar yang sedang ditempati oleh penghuni, "
     "meliputi nomor kamar, tipe, fasilitas, harga sewa, serta tanggal mulai sewa. "
     "Informasi ini memberikan kejelasan kepada penghuni mengenai kamar dan fasilitas "
     "yang menjadi haknya."),

    ("11_tagihan_saya", "Halaman Tagihan Saya",
     "Halaman Tagihan Saya menampilkan daftar tagihan sewa milik penghuni beserta "
     "statusnya (lunas atau belum lunas). Untuk tagihan yang belum dibayar, tersedia "
     "tombol \"Bayar\" yang akan mengarahkan penghuni ke halaman pembayaran. Penghuni "
     "dapat memantau riwayat dan kewajiban pembayaran melalui halaman ini."),

    ("12_form_bayar", "Form Pembayaran dan Unggah Bukti Transfer",
     "Form pembayaran digunakan penghuni untuk melakukan pembayaran tagihan sewa. "
     "Penghuni mengisi jumlah bayar, tanggal transfer, metode pembayaran, serta "
     "mengunggah bukti transfer dalam format JPG, PNG, atau PDF (maksimal 2 MB). Sistem "
     "menampilkan pratinjau (preview) gambar sebelum diunggah dan menampilkan informasi "
     "rekening tujuan transfer. Setelah disubmit, status pembayaran menjadi menunggu "
     "konfirmasi pemilik."),

    ("13_ajukan_komplain", "Form Ajukan Komplain",
     "Form ini digunakan penghuni untuk mengajukan komplain terkait fasilitas, "
     "kebersihan, keamanan, atau hal lain. Penghuni memilih kategori komplain, mengisi "
     "judul dan deskripsi masalah, serta dapat melampirkan foto sebagai bukti. Komplain "
     "yang diajukan akan diteruskan ke pemilik untuk ditindaklanjuti."),

    ("14_profil", "Halaman Profil",
     "Halaman Profil menampilkan data diri pengguna yang sedang login. Pengguna dapat "
     "melihat dan memperbarui informasi profil seperti nama, email, dan nomor telepon. "
     "Fitur ini tersedia untuk seluruh peran pengguna."),

    ("15_ubah_password", "Halaman Ubah Password",
     "Halaman Ubah Password memungkinkan pengguna mengganti kata sandi akun. Pengguna "
     "harus memasukkan kata sandi lama, kata sandi baru, dan konfirmasi kata sandi baru. "
     "Sistem memvalidasi kebenaran kata sandi lama sebelum perubahan disimpan untuk "
     "menjaga keamanan akun."),
]


def build_bab4(doc):
    """Membuat BAB 4 Hasil."""
    # Judul BAB
    add_para(doc, "BAB IV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_para(doc, "HASIL DAN PEMBAHASAN", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=12)

    # Pengantar
    add_para(doc, "4.1  Gambaran Umum Aplikasi", bold=True, space_after=6)
    intro = (
        'Aplikasi Sistem Manajemen Kost "KostKu" merupakan aplikasi web yang dibangun '
        "menggunakan framework Laravel. Aplikasi ini dikembangkan berdasarkan hasil "
        "analisis dan perancangan pada Projek 1. Aplikasi memiliki tiga peran pengguna, "
        "yaitu Admin, Pemilik Kost, dan Penghuni, dengan hak akses yang berbeda-beda. "
        "Secara umum, aplikasi ini menyediakan fitur pengelolaan kamar, penghuni, "
        "pembayaran sewa, serta penanganan komplain. Pada bab ini disajikan hasil "
        "implementasi aplikasi dalam bentuk tangkapan layar (screenshot) beserta "
        "penjelasan dari masing-masing halaman."
    )
    pp = add_para(doc, intro, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)
    pp.paragraph_format.first_line_indent = Cm(1)

    add_para(doc, "4.2  Hasil Implementasi", bold=True, space_after=6)

    for i, (fname, judul, desc) in enumerate(HASIL, start=1):
        # Subjudul
        add_para(doc, f"4.2.{i}  {judul}", bold=True, space_after=6)
        # Screenshot + caption
        add_screenshot(doc, fname, i, judul)
        # Deskripsi
        p = add_para(doc, desc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)
        p.paragraph_format.first_line_indent = Cm(1)


def main():
    doc = Document()
    set_base_style(doc)
    set_margins(doc.sections[0])

    # Cover
    build_cover(doc)

    # Page break ke BAB 4
    doc.add_page_break()
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(new_sec)

    build_bab4(doc)

    doc.save(OUT)
    print(f"BERHASIL: Laporan disimpan di -> {OUT}")

    # Hitung screenshot yang ditemukan
    found = 0
    for fname, _, _ in HASIL:
        for ext in (".png", ".jpg", ".jpeg"):
            if os.path.exists(os.path.join(SHOT, fname + ext)):
                found += 1
                break
    print(f"Screenshot ditemukan: {found}/{len(HASIL)}")
    if found < len(HASIL):
        print("Sisanya memakai placeholder. Tambahkan gambar lalu jalankan ulang script.")


if __name__ == "__main__":
    main()
